from contextlib import contextmanager
import docx
from typing import Generator

from src.core.templates.iface import TemplatesFileStorageABC
from src.core.templates.types import Template


class TemplateFileService:

    templates_storage: TemplatesFileStorageABC

    def __init__(self, templates_storage: TemplatesFileStorageABC):
        self.templates_storage = templates_storage

    @contextmanager
    def __open_docx(self, filename: str) -> Generator[docx.Document, None, None]:
        with self.templates_storage.open_template_file(filename) as file:
            yield docx.Document(file)

    def extract_text(self, template: Template) -> str:
        with self.__open_docx(template.storage_filename) as doc:
            chunks = []
            for p in doc.paragraphs:
                if p.text.strip():
                    chunks.append(p.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        txt = cell.text.strip()
                        if txt:
                            chunks.append(txt)

            return "\n".join(chunks)
